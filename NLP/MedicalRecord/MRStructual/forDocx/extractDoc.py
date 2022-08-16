import docx
from docx import Document
import os
from win32com import client as wc
# word = wc.Dispatch('Word.Application')
word = wc.gencache.EnsureDispatch('kwps.application')

from Lib.Utils import Utils
utils = Utils()

def process_table(table):
    """
    合并表格内容
    """
    cnts = []
    for row in table.rows:
        for cell in row.cells:
            txt = cell.text.strip()
            if len(txt) > 0:
                if len(cnts) == 0 or cnts[-1] != txt:
                    cnts.append(txt)

    return cnts


def process_header_line(lines):
    """
    生成头部第一行
    """
    items = []
    for idx in range(len(lines) - 1, -1, -1):
        line = lines[idx]
        if '：' in line:
            str = lines[idx].split("：")[1]
        elif ':' in line:
            str = lines[idx].split(":")[1]
        else:
            str = line
        items.append(str)

    return '||'.join(items)


def docToDocx(file_path):
    """
    doc转docx
    """
    print('转换文件：%s' % file_path)
    if file_path.endswith('.docx'):
        return file_path
    else:
        doc = word.Documents.Open(file_path)
        docxNamePath = file_path.replace('.doc', '.docx')
        # doc.SaveAs(docxNamePath, 12, False, "", True, "", False, False, False, False)
        doc.SaveAs2(docxNamePath, 12)
        print('转换完成!')
        doc.Close()
        os.remove(file_path)
        return docxNamePath


def write_to_txt(lines, file_path):
    """
    将结果写入txt
    """
    with open(file_path, 'a+') as f:
        for line in lines:
            f.write(line + '\n')
        f.write('\n\n')


# 入院记录
def process_ryjl(file_path):
    """
    处理入院记录、出院记录
    """
    doc = Document(file_path)
    header = doc.sections[0].header
    cnts = []
    # 添加页眉部分内容
    header_line = process_header_line(process_table(header.tables[0]))
    print('header line: %s' % header_line)
    cnts.append(header_line)

    # 添加第一个表格
    if len(doc.tables) > 0:
        cnts.extend(process_table(doc.tables[0]))

    # 添加其余内容
    for p in doc.paragraphs:
        if len(p.text.strip()) > 0:
            cnts.append(p.text.strip())

    return cnts


# 病程记录
def process_bcjl(file_path):
    """
    病程记录
    """
    doc = Document(file_path)
    header = doc.sections[0].header
    cnts_arr = []
    cnts_scbc, cnts_rcbc = [], []   # 首次病程，日常病程

    # 添加页眉部分内容
    header_line = process_header_line(process_table(header.tables[0]))
    cnts_scbc.append(header_line)
    print('header line: %s' % header_line)

    # 每一个病程放一个数组
    dt_cnt = 0
    for p in doc.paragraphs:
        if len(p.text.strip()) > 0:
            txt = p.text.strip()
            dt = utils.format_date(txt)
            if len(dt) > 0 and txt.startswith(dt[:4]):
                dt_cnt = dt_cnt + 1
                if dt_cnt >= 2:
                    cnts_rcbc.append('')
                    cnts_rcbc.append('')
                    cnts_rcbc.append(header_line)

            if dt_cnt < 2:
                cnts_scbc.append(txt)
            else:
                cnts_rcbc.append(txt)

    return cnts_scbc, cnts_rcbc



if __name__ == "__main__":
    in_folder = r'D:\项目资料\病历辅助诊断\人机大赛数据集'
    out_folder = r'D:\projects\NLP\MedicalRecord\FeatureExtraction\data\202207人机大赛'
    ry_txt = os.path.join(out_folder, '入院记录.txt')
    cy_txt = os.path.join(out_folder, '出院记录.txt')
    sc_txt = os.path.join(out_folder, '首次病程.txt')
    rc_txt = os.path.join(out_folder, '日常病程.txt')

    for folder_1 in os.listdir(in_folder):
        print('folder_1: %s' % folder_1)
        folder_1_path = os.path.join(in_folder, folder_1)
        if os.path.isdir(folder_1_path):
            print(os.listdir(folder_1_path))
            for folder_2 in os.listdir(folder_1_path):
                print('folder_2: %s' % folder_2)
                folder_2_path = os.path.join(folder_1_path, folder_2)
                if os.path.isdir(folder_2_path):
                    print(os.listdir(folder_2_path))
                    for file in os.listdir(folder_2_path):
                        print('file: %s' % file)
                        file_path = os.path.join(folder_2_path, file)
                        if os.path.isfile(file_path) and not file.startswith('~') and (file_path.endswith('.doc') or file_path.endswith('.docx')):
                            # file_path = r'D:\项目资料\病历辅助诊断\人机大赛数据集\急性胰腺炎30\曾令兰\入院记录.docx'
                            # file_path = r'D:\项目资料\病历辅助诊断\人机大赛数据集\肠道穿孔14\0002851213\首次病程.doc'
                            file_path = docToDocx(file_path)
                            file_name = file_path.split('\\')[-1]
                            if '入院' in file_name:
                                cnts = process_ryjl(file_path)
                                write_to_txt(cnts, ry_txt)
                            elif '出院' in file_name:
                                cnts = process_ryjl(file_path)
                                write_to_txt(cnts, cy_txt)
                            elif '首' in file_name and '程' in file_name:
                                cnts_scbc, cnts_rcbc = process_bcjl(file_path)
                                write_to_txt(cnts_scbc, sc_txt)
                                write_to_txt(cnts_rcbc, rc_txt)





#
